import os
import re
import json
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn


# ─────────────────────────────────────────────
# GROUP 1 PATTERNS
# ─────────────────────────────────────────────

MRN_PATTERNS = [
    re.compile(
        r'\b(MRN|Medical\s*Record\s*(Number)?|Medical\s*Record\s*No\.?'
        r'|Record\s*#|MR|Patient\s*ID)[:\s#]*([A-Z0-9\-]{4,12})\b',
        re.IGNORECASE
    ),
]

DOB_LABEL_PATTERNS = [
    re.compile(
        r'\b(DOB|D\.O\.B\.?|Date\s*of\s*Birth|Birth\s*Date|Birthdate|Born)[:\s]*'
        r'('
        r'\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}'
        r'|'
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}'
        r'|'
        r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}'
        r')',
        re.IGNORECASE
    ),
]

SSN_PATTERN = re.compile(r'\b\d{3}-\d{2}-\d{4}\b')

PHONE_PATTERN = re.compile(
    r'(?<!\d)'
    r'('
    r'\+?1?[\s\-\.]?\(?\d{3}\)?[\s\-\.]\d{3}[\s\-\.]\d{4}(?:\s*(?:x|ext\.?)\s*\d{1,6})?'
    r'|'
    r'(?:001|\+1)[\s\-]?\d{3}[\s\-]\d{3}[\s\-]\d{4}(?:\s*(?:x|ext\.?)\s*\d{1,6})?'
    r'|'
    r'\(\d{3}\)\s*\d{3}[\s\-]\d{4}(?:\s*(?:x|ext\.?)\s*\d{1,6})?'
    r')'
    r'(?!\d)',
    re.IGNORECASE
)

FAX_PATTERN = re.compile(
    r'\b(?:Fax|FAX|Fax\s*#|Fax\s*Number)[:\s]*'
    r'[\+]?[\d\s\-\(\)\.]{7,20}\b',
    re.IGNORECASE
)

EMAIL_PATTERN = re.compile(
    r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
)

IP_PATTERN = re.compile(
    r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b'
)

URL_PATTERN = re.compile(
    r'https?://[^\s<>"\']+|www\.[^\s<>"\']+',
    re.IGNORECASE
)


# ─────────────────────────────────────────────
# GROUP 2 PATTERNS
# ─────────────────────────────────────────────

# All dates more specific than year-only.
# Validated ranges prevent IP addresses matching as dates.
ALL_DATE_PATTERN = re.compile(
    r'\b(?:'
    r'(?:0?[1-9]|1[0-2])[\/\-\.](?:0?[1-9]|[12]\d|3[01])[\/\-\.]\d{2,4}'
    r'|(?:January|February|March|April|May|June|July|August'
    r'|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
    r'|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s+\d{4}'
    r'|(?:January|February|March|April|May|June|July|August'
    r'|September|October|November|December)\s+\d{4}'
    r'|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{4}'
    r'|\d{4}-(?:0[1-9]|1[0-2])-(?:0[1-9]|[12]\d|3[01])'
    r')\b',
    re.IGNORECASE
)

# Ages over 89
AGE_OVER_89_PATTERN = re.compile(
    r'\b(9\d|1[0-9]{2})[- ]year[- ]old\b'
    r'|\b(?:Age|AGE)[:\s]+(9\d|1[0-9]{2})\b',
    re.IGNORECASE
)

# ZIP codes — labeled or address context (State abbreviation + 5 digits)
ZIP_PATTERN = re.compile(
    r'\b(?:ZIP|Zip\s*Code|Postal\s*Code)[:\s]*\d{5}(?:-\d{4})?'
    r'|\b[A-Z]{2}\s+\d{5}(?:-\d{4})?\b'
)

# NPI — 10-digit National Provider Identifier
NPI_PATTERN = re.compile(
    r'(\b(?:NPI|National\s*Provider\s*(?:Identifier|ID|Number)?)[:\s#]*)(\d{10})\b',
    re.IGNORECASE
)

# Account numbers — labeled only
ACCOUNT_PATTERN = re.compile(
    r'(\b(?:Account\s*(?:Number|No\.?|#)|Acct\.?\s*(?:Number|No\.?|#)?)[:\s#]*)([A-Z0-9\-]{4,17})\b',
    re.IGNORECASE
)

# Beneficiary / Member / Policy / Group / Insurance IDs — labeled only.
# Requires colon or # separator to prevent facility names with "Group" matching.
BENEFICIARY_PATTERN = re.compile(
    r'(\b(?:Beneficiary\s*(?:ID|Number|No\.?)?|Member\s*(?:ID|Number|No\.?)?'
    r'|Subscriber\s*(?:ID|Number|No\.?)?|Policy\s*(?:Number|No\.?)?'
    r'|Group\s*(?:Number|No\.?)?|Plan\s*(?:ID|Number)'
    r'|Insurance\s*(?:ID|Number|No\.?)?)\s*[:#]\s*)([A-Z0-9\-]{4,20})\b',
    re.IGNORECASE
)

# Address — labeled only. Negative lookbehind prevents "IP address:" matching.
ADDRESS_PATTERN = re.compile(
    r'(?<![Ii][Pp] )(\b(?:Patient\s*)?Address[:\s]+)([^\n\|]{10,120})',
    re.IGNORECASE
)

# License / DEA / Certificate numbers — labeled only
LICENSE_PATTERN = re.compile(
    r'(\b(?:(?:Medical\s*)?License\s*(?:Number|No\.?|#)?'
    r'|DEA\s*(?:Number|No\.?|#)?'
    r'|(?:State|Professional)\s*License)[:\s#]*)([A-Z0-9\-]{4,15})\b',
    re.IGNORECASE
)

# Device / Serial / UDI numbers — labeled only
DEVICE_PATTERN = re.compile(
    r'(\b(?:(?:Device|Serial|Implant)\s*(?:ID|Number|No\.?|#)|UDI|Serial\s*#)[:\s#]*)([A-Z0-9\-\/]{4,20})\b',
    re.IGNORECASE
)

# VINs — labeled only (17-char alphanumeric, no I/O/Q per ISO 3779)
VIN_PATTERN = re.compile(
    r'(\b(?:VIN|Vehicle\s*Identification\s*(?:Number)?|License\s*Plate\s*(?:Number)?)[:\s#]*)([A-HJ-NPR-Z0-9]{17})\b',
    re.IGNORECASE
)


# ─────────────────────────────────────────────
# NAME LABEL PATTERNS
# ─────────────────────────────────────────────

NAME_LABEL_PATTERNS = [
    re.compile(
        r'\b(Patient[\'\s]*s?\s*Name|Patient\s*Name|Full\s*Name|Patient|Name)[:\s]+'
        r'([A-Z][a-zA-Z\-\']+(?:,\s*[A-Z][a-zA-Z\-\']+)?(?:\s+[A-Z][a-zA-Z\-\'\.]+){0,3})'
        r'(?=\s*$|\s*\n|\s*\||\s+(?:DOB|MRN|SSN|Phone|Address|Age|DATE|FILE|is\s+a\b|was\b|presents\b))',
        re.IGNORECASE
    ),
    re.compile(
        r'\bPatient\s+'
        r'([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\'\.]+){0,3})'
        r'(?=\s+(?:is|was|has|had|presents|presented|denies|reports|states|called|reached)\b)',
        re.IGNORECASE
    ),
    re.compile(
        r'\b(Mr|Mrs|Ms|Miss)\.?\s+'
        r'('
        r'[A-Z][A-Za-z\-\']+,\s+[A-Z][A-Za-z\-\']+(?:\s+[A-Z]\.?)?'
        r'|'
        r'[A-Z][a-zA-Z\-\']+(?:\s+[A-Z]\.?)?(?:\s+[A-Z][a-zA-Z\-\']+){0,2}'
        r')',
        re.IGNORECASE
    ),
    re.compile(
        r'\bRe:\s+'
        r'([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+){1,3})'
        r'(?=\s*,|\s+DOB|\s+MRN)',
        re.IGNORECASE
    ),
    re.compile(
        r'\bpatient,?\s+'
        r'([A-Z][A-Za-z\-\']+,\s+[A-Z][A-Za-z\-\']+(?:\s+[A-Z]\.?)?)'
        r'(?=\s*,|\s+DOB|\s+MRN|\s+\()',
        re.IGNORECASE
    ),
    re.compile(
        r'\bpatient,?\s+'
        r'([A-Z][A-Za-z\-\']+\s+[A-Z][A-Za-z\-\']+(?:\s+[A-Z]\.?)?)'
        r'(?=\s*,|\s+DOB|\s+MRN|\s+\()',
        re.IGNORECASE
    ),
    re.compile(
        r'\b([A-Z][A-Za-z\-\']+),\s+([A-Z][A-Za-z\-\']+(?:\s+[A-Z]\.?)?)'
        r'(?=[^.]{0,80}(?:DOB|D\.O\.B|Date\s+of\s+Birth|MRN|Medical\s+Record))',
        re.IGNORECASE
    ),
]


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

NON_NAME_WORDS = {
    'demographics', 'information', 'details', 'summary', 'report',
    'note', 'notes', 'record', 'data', 'history', 'profile',
    'unknown', 'confidential', 'patient', 'laboratory', 'discharge',
    'referral', 'clinical', 'medical', 'intake', 'admission',
    'letter', 'evaluation', 'consultation', 'assessment', 'follow', 'update',
    'the', 'and', 'for', 'with', 'date', 'birth',
    'none', 'n/a', 'na', 'mr', 'mrs', 'ms', 'miss', 'dr'
}

SUFFIXES = {
    'i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii',
    'jr', 'jr.', 'sr', 'sr.', 'esq', 'esq.',
    'md', 'md.', 'do', 'do.', 'phd', 'phd.'
}


def _normalize_name(name: str) -> str:
    name = name.strip()
    comma_match = re.match(
        r"^([A-Za-z][A-Za-z\-\']*(?:\s+[A-Za-z][A-Za-z\-\']*)*)"
        r",\s*([A-Za-z][A-Za-z\-\']*)((?:\s+[A-Za-z]\.?)*)?$",
        name
    )
    if comma_match:
        last  = comma_match.group(1).strip().title()
        first = comma_match.group(2).strip().title()
        mid   = (comma_match.group(3) or "").strip().title()
        return f"{first} {mid} {last}".strip() if mid else f"{first} {last}"
    if name.isupper():
        return name.title()
    return name


def _clean_name(name: str) -> str:
    return re.sub(r'[\s.,;:]+$', '', name.strip())


def _split_name_parts(name: str):
    parts = name.strip().split()
    if len(parts) < 2:
        return None, [], None
    while parts and parts[-1].lower() in SUFFIXES:
        parts.pop()
    if len(parts) < 2:
        return None, [], None
    return parts[0], parts[1:-1], parts[-1]


def _build_name_variants(original_name: str) -> list[str]:
    original_name = original_name.strip()
    if not original_name:
        return []

    normalized = _normalize_name(original_name)
    first, middles, last = _split_name_parts(normalized)
    if first is None:
        return []

    variants = set()
    variants.add(f"{first} {' '.join(middles)} {last}".strip() if middles else f"{first} {last}")
    variants.add(f"{first} {last}")
    variants.add(f"{last}, {first}")
    variants.add(f"{last.upper()}, {first.upper()}")
    variants.add(f"{first.upper()} {last.upper()}")

    if middles:
        variants.add(f"{last}, {first} {' '.join(middles)}")
        variants.add(f"{first.upper()} {' '.join(m.upper() for m in middles)} {last.upper()}")
        mid_initial = middles[0][0]
        variants.add(f"{first} {mid_initial}")
        variants.add(f"{first} {mid_initial}.")

    if len(last) >= 5:
        variants.add(last)
        variants.add(last.upper())
    if len(first) >= 5:
        variants.add(first)
        variants.add(first.upper())

    if '-' in last:
        for part in last.split('-'):
            if len(part) >= 5:
                variants.add(part)

    variants.add(original_name)

    filtered = sorted(
        {v for v in variants
         if len(v.strip()) >= 2 and v.strip().lower() not in NON_NAME_WORDS},
        key=len, reverse=True
    )
    seen, unique = set(), []
    for v in filtered:
        k = v.strip().lower()
        if k not in seen:
            seen.add(k)
            unique.append(v.strip())
    return unique


def _extract_name_from_page1(doc: Document) -> str | None:

    def is_valid_name(name: str) -> bool:
        if not name or len(name.strip()) < 2:
            return False
        parts = name.strip().split()
        if len(parts) == 1 and parts[0].lower() in NON_NAME_WORDS:
            return False
        if not name[0].isupper():
            return False
        if any(p.lower() in NON_NAME_WORDS for p in parts):
            return False
        if not re.match(r"^[A-Za-z\s\-\'\.]+$", name):
            return False
        return True

    # Priority 0: Bold standalone paragraph
    for para in doc.paragraphs[:5]:
        text   = para.text.strip()
        is_bold = para.runs and any(run.bold for run in para.runs)
        if is_bold and re.match(r'^[A-Z][a-zA-Z\-\']+([\s][A-Z][a-zA-Z\-\']+){1,3}$', text):
            if is_valid_name(text):
                return _clean_name(_normalize_name(text))

    # Priority 1: Re: line
    for para in doc.paragraphs[:30]:
        m = NAME_LABEL_PATTERNS[3].search(para.text.strip())
        if m:
            name = _clean_name(_normalize_name(m.group(1).strip()))
            if is_valid_name(name):
                return name

    # Priority 2: Table label → value
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        for row in rows:
            cells = [c.text.strip() for c in row.cells]
            for i, cell in enumerate(cells):
                if re.match(r'^(Patient\s*Name|Full\s*Name|Name)$', cell, re.IGNORECASE):
                    if i + 1 < len(cells) and cells[i + 1]:
                        candidate = _normalize_name(cells[i + 1].strip())
                        if is_valid_name(candidate):
                            return _clean_name(candidate)
        if len(rows) >= 2:
            header_cells = [c.text.strip() for c in rows[0].cells]
            for col_idx, header in enumerate(header_cells):
                if re.match(r'^(Patient\s*Name|Full\s*Name|Name)$', header, re.IGNORECASE):
                    value = rows[1].cells[col_idx].text.strip()
                    if value:
                        candidate = _normalize_name(value)
                        if is_valid_name(candidate):
                            return _clean_name(candidate)
        break

    # Priority 3: Paragraph patterns
    for para in doc.paragraphs[:30]:
        text = para.text.strip()
        if not text:
            continue
        for idx, pat in enumerate(NAME_LABEL_PATTERNS):
            m = pat.search(text)
            if not m:
                continue
            grp = 2 if idx in (0, 2) else (1 if idx in (1, 5) else
                  (1 if idx in (3, 4) else None))
            if grp is None:
                if idx == 6:
                    raw  = f"{m.group(1)}, {m.group(2)}"
                    name = _clean_name(_normalize_name(raw))
                    if is_valid_name(name):
                        return name
                continue
            name = _clean_name(_normalize_name(m.group(grp).strip()))
            if is_valid_name(name):
                return name

    return None


# ─────────────────────────────────────────────
# MAIN REDACTION FUNCTION
# ─────────────────────────────────────────────

def redact_docx(input_path: str, document_id: str = "unknown",
                output_dir: str = "output_docs") -> dict:
    """
    Performs targeted in-place redaction of a DOCX file.
    Redacts Group 1 + Group 2 Safe Harbor identifiers.
    Also redacts headers and footers.
    Preserves all DOCX formatting, fonts, tables, and structure.
    """
    doc = Document(input_path)
    counts = {
        # Group 1
        "PATIENT_NAME": 0, "DATE_OF_BIRTH": 0, "MRN": 0,
        "SSN": 0, "PHONE": 0, "FAX": 0, "EMAIL": 0,
        "IP_ADDRESS": 0, "URL": 0,
        # Group 2
        "DATE": 0, "AGE": 0, "ZIP": 0, "NPI": 0,
        "ADDRESS": 0, "ACCOUNT_NUMBER": 0, "BENEFICIARY_ID": 0,
        "LICENSE_NUMBER": 0, "DEVICE_ID": 0, "VIN": 0,
    }

    # Step 1: Discover patient name
    patient_name  = _extract_name_from_page1(doc)
    name_variants = _build_name_variants(patient_name) if patient_name else []

    # Step 2: Build ordered replacement list
    # Each entry: (compiled_pattern, replacement_string, entity_type_key)
    tagged = []

    # ── Group 1 ──
    tagged.append((
        re.compile(
            r'\b(MRN|Medical\s*Record\s*(?:Number)?|Medical\s*Record\s*No\.?'
            r'|Record\s*#|Patient\s*ID)([:\s#]*)(\d{5,12})\b',
            re.IGNORECASE
        ),
        r'\1\2[MRN]', "MRN"
    ))
    tagged.append((
        re.compile(
            r'(\b(?:DOB|D\.O\.B\.?|Date\s*of\s*Birth|Birth\s*Date|Birthdate|Born)[:\s]*)'
            r'(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}'
            r'|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}'
            r'|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})',
            re.IGNORECASE
        ),
        r'\1[DATE OF BIRTH]', "DATE_OF_BIRTH"
    ))
    tagged.append((SSN_PATTERN,   '[SSN]',   "SSN"))
    tagged.append((URL_PATTERN,   '[URL]',   "URL"))
    tagged.append((FAX_PATTERN,   '[FAX]',   "FAX"))
    tagged.append((PHONE_PATTERN, '[PHONE]', "PHONE"))
    tagged.append((IP_PATTERN,    '[IP ADDRESS]', "IP_ADDRESS"))
    tagged.append((EMAIL_PATTERN, '[EMAIL]', "EMAIL"))

    # ── Group 2 ──
    # DOB runs first above so its dates are already replaced as [DATE OF BIRTH]
    # before ALL_DATE scans — no double-redaction possible in DOCX.
    # ADDRESS runs before ZIP so the full address is replaced first;
    # ZIP won't find state+zip text that's already been substituted.
    tagged.append((ALL_DATE_PATTERN,    '[DATE]',           "DATE"))
    tagged.append((AGE_OVER_89_PATTERN, '[AGE]',            "AGE"))
    tagged.append((ADDRESS_PATTERN,     r'\1[ADDRESS]',     "ADDRESS"))
    tagged.append((ZIP_PATTERN,         '[ZIP]',            "ZIP"))
    tagged.append((NPI_PATTERN,         r'\1[NPI]',             "NPI"))
    tagged.append((ACCOUNT_PATTERN,     r'\1[ACCOUNT NUMBER]',  "ACCOUNT_NUMBER"))
    tagged.append((BENEFICIARY_PATTERN, r'\1[BENEFICIARY ID]',  "BENEFICIARY_ID"))
    tagged.append((LICENSE_PATTERN,     r'\1[LICENSE NUMBER]',  "LICENSE_NUMBER"))
    tagged.append((DEVICE_PATTERN,      r'\1[DEVICE ID]',       "DEVICE_ID"))
    tagged.append((VIN_PATTERN,         r'\1[VIN]',             "VIN"))

    # Name variants — longest first
    for variant in name_variants:
        tagged.append((
            re.compile(r'\b' + re.escape(variant) + r'\b', re.IGNORECASE),
            '[PATIENT NAME]', "PATIENT_NAME"
        ))

    def _apply(text: str) -> tuple[str, dict]:
        """Apply all replacements to a string, return (new_text, delta_counts)."""
        delta = {k: 0 for k in counts}
        for pattern, replacement, entity in tagged:
            new_text, n = re.subn(pattern, replacement, text)
            delta[entity] += n
            text = new_text
        return text, delta

    def _process_para(para):
        full = "".join(r.text for r in para.runs)
        if not full.strip():
            return
        new_text, delta = _apply(full)
        if new_text != full and para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
        for k, v in delta.items():
            counts[k] += v

    # Step 3: Body paragraphs
    for para in doc.paragraphs:
        _process_para(para)

    # Step 4: Tables
    for table in doc.tables:
        seen_ids = set()
        for row in table.rows:
            cells = row.cells
            i = 0
            while i < len(cells):
                cell    = cells[i]
                cell_id = id(cell._tc)
                if cell_id in seen_ids:
                    i += 1
                    continue
                seen_ids.add(cell_id)
                label = cell.text.strip().lower()

                # Label-value pairs — redact the next cell wholesale
                if re.match(r'^(mrn|medical record(?:\s*number)?|record #|patient id)$', label):
                    if i + 1 < len(cells):
                        val = cells[i + 1]
                        seen_ids.add(id(val._tc))
                        for p in val.paragraphs:
                            if p.text.strip() and p.runs:
                                p.runs[0].text = '[MRN]'
                                for r in p.runs[1:]: r.text = ''
                                counts["MRN"] += 1
                    i += 2; continue

                elif re.match(r'^(dob|date of birth|birth date|birthdate|born)$', label):
                    if i + 1 < len(cells):
                        val = cells[i + 1]
                        seen_ids.add(id(val._tc))
                        for p in val.paragraphs:
                            if p.text.strip() and p.runs:
                                p.runs[0].text = '[DATE OF BIRTH]'
                                for r in p.runs[1:]: r.text = ''
                                counts["DATE_OF_BIRTH"] += 1
                    i += 2; continue

                elif re.match(r'^(patient\s*name|full\s*name)$', label):
                    if i + 1 < len(cells):
                        val = cells[i + 1]
                        seen_ids.add(id(val._tc))
                        for p in val.paragraphs:
                            if p.text.strip() and p.runs:
                                p.runs[0].text = '[PATIENT NAME]'
                                for r in p.runs[1:]: r.text = ''
                                counts["PATIENT_NAME"] += 1
                    i += 2; continue

                # All other cells — apply full replacement pipeline
                for cell_para in cell.paragraphs:
                    _process_para(cell_para)
                i += 1

    # Step 5: Headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is None:
                continue
            for para in hf.paragraphs:
                _process_para(para)

    # Step 6: Save
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(input_path))[0].replace("_converted", "")
    output_path = os.path.join(output_dir, f"REDACTED_{base_name}.docx")
    doc.save(output_path)

    # Step 7: Audit log
    total  = sum(counts.values())
    result = {
        "document_id":             document_id,
        "timestamp":               datetime.utcnow().isoformat() + "Z",
        "output_path":             output_path,
        "patient_name_discovered": patient_name or "not found",
        "name_variants_searched":  name_variants,
        "entity_counts":           counts,
        "total_redactions":        total,
    }

    log_dir  = "logs"
    os.makedirs(log_dir, exist_ok=True)
    ts       = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    log_path = os.path.join(log_dir, f"{ts}_{document_id}_redaction_log.json")
    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2)
    result["log_path"] = log_path

    return result