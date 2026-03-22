import pdfplumber
import docx
import os


def extract_text(file_path: str) -> str:
    """
    Accepts a path to a PDF or DOCX file.
    Returns the full extracted text as a single string.
    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def _extract_from_pdf(file_path: str) -> str:
    """
    Extracts text from all pages of a PDF.
    Joins pages with a newline separator.
    """
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
            else:
                # Flag blank/scanned pages — OCR not handled here yet
                pages.append(f"[PAGE {i+1}: NO EXTRACTABLE TEXT — MAY BE SCANNED]")
    return "\n\n".join(pages)


def _extract_from_docx(file_path: str) -> str:
    """
    Extracts text from all paragraphs in a DOCX file.
    Tables are extracted cell by cell.
    """
    doc = docx.Document(file_path)
    content = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                content.append(row_text)

    return "\n".join(content)