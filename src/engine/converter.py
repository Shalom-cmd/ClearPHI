import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import io
import shutil

import platform

if platform.system() == "Windows":
    tesseract_path = shutil.which("tesseract") or r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    tesseract_path = shutil.which("tesseract") or "/opt/homebrew/bin/tesseract"

pytesseract.pytesseract.tesseract_cmd = tesseract_path


def pdf_to_docx(pdf_path: str) -> str:
    """
    Converts a PDF to DOCX using PyMuPDF only.
    - Digital PDFs: text extracted directly via PyMuPDF blocks
    - Scanned PDFs: auto-detected, rendered at 300 DPI, OCR via Tesseract
    
    pdf2docx has been removed entirely — it crashes Gunicorn workers on
    complex nested tables via infinite recursion in Cell.make_docx().
    """
    base = os.path.splitext(pdf_path)[0]
    docx_path = base + "_converted.docx"

    if _is_scanned(pdf_path):
        return _ocr_pdf_to_docx(pdf_path, docx_path)
    else:
        return _digital_pdf_to_docx(pdf_path, docx_path)


def _is_scanned(pdf_path: str) -> bool:
    """
    Returns True if the PDF has little or no extractable text —
    i.e. it is image-only and needs OCR.
    """
    try:
        pdf = fitz.open(pdf_path)
        total_text = "".join(page.get_text() for page in pdf)
        pdf.close()
        return len(total_text.strip()) < 100
    except Exception:
        return True


def _digital_pdf_to_docx(pdf_path: str, output_path: str) -> str:
    """
    Extract text from a digital PDF using PyMuPDF block extraction.
    Preserves paragraph structure as best as possible.
    Builds a clean DOCX from the extracted text blocks.
    """
    pdf = fitz.open(pdf_path)
    doc = DocxDocument()

    for page_num, page in enumerate(pdf):
        blocks = page.get_text("blocks")  # returns list of (x0,y0,x1,y1,text,block_no,block_type)
        # Sort top-to-bottom, left-to-right
        blocks.sort(key=lambda b: (round(b[1] / 10), b[0]))

        for block in blocks:
            text = block[4].strip()
            if not text:
                continue
            # Each block becomes a paragraph
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        # Page separator (except last page)
        if page_num < len(pdf) - 1:
            doc.add_paragraph("─" * 40)

    pdf.close()
    doc.save(output_path)
    return output_path


def _ocr_pdf_to_docx(pdf_path: str, output_path: str) -> str:
    """
    OCR pipeline for scanned/image-only PDFs.
    Renders each page at 300 DPI and runs Tesseract on it.
    Builds a clean DOCX from the extracted text.
    """
    pdf = fitz.open(pdf_path)
    doc = DocxDocument()

    for page_num, page in enumerate(pdf):
        # Render at 300 DPI for good OCR accuracy
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)

        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        ocr_text = pytesseract.image_to_string(img, lang='eng')

        for line in ocr_text.split('\n'):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

        if page_num < len(pdf) - 1:
            doc.add_paragraph("─" * 40)

    pdf.close()
    doc.save(output_path)
    return output_path


def is_scanned_pdf(pdf_path: str) -> bool:
    """
    Public helper — returns True if the PDF appears to be scanned.
    Useful for logging and debugging.
    """
    return _is_scanned(pdf_path)