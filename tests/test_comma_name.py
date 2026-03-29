import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.engine.redactor import redact_docx, _extract_name_from_page1, _normalize_name, _build_name_variants
from docx import Document

# Test normalize function directly
print("=== Normalize Tests ===")
print(_normalize_name("MAKENA, SHALOM"))
print(_normalize_name("DOE, REGINA"))
print(_normalize_name("JOHN SMITH"))
print(_normalize_name("Eric Robert Jamison"))

# Create a tiny test doc with PATIENT: MAKENA, SHALOM format
print("\n=== Document Test ===")
doc = Document()
doc.add_paragraph("PATIENT: MAKENA, SHALOM")
doc.add_paragraph("DOB: 03/22/1975")
doc.add_paragraph("MRN: 83055723")
doc.add_paragraph("Patient was seen today for follow up.")
test_path = "input_docs/test_comma.docx"
os.makedirs("input_docs", exist_ok=True)
doc.save(test_path)

result = redact_docx(test_path, "comma_test")
print("Name found:", result['patient_name_discovered'])
print("Variants:", result['name_variants_searched'])
print("Counts:", result['entity_counts'])

# Check output
out_doc = Document(result['output_path'])
for para in out_doc.paragraphs:
    if para.text.strip():
        print(para.text)
