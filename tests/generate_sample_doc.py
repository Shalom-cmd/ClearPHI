import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

doc = Document()

doc.add_heading("Occidental Medical Center — Patient Clinical Note", 0)

doc.add_paragraph("Date of Visit: 03/15/2026")
doc.add_paragraph("Prepared by: Dr. Robert Langley, MD | NPI: NPI:1234567890")
doc.add_paragraph("Fax: Fax: 555-234-5678")

doc.add_heading("Patient Information", level=1)

table = doc.add_table(rows=6, cols=2)
table.style = "Table Grid"

rows = [
    ("Patient Name", "Maria Elena Gonzalez"),
    ("Date of Birth", "07/04/1981"),
    ("SSN", "456-78-9012"),
    ("MRN", "MRN#789012"),
    ("Insurance Member ID", "Member ID: XYZ987654"),
    ("Address", "456 Oak Avenue, Dallas, TX 75201"),
]

for i, (label, value) in enumerate(rows):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_heading("Clinical Notes", level=1)
doc.add_paragraph(
    "Patient Maria Elena Gonzalez is a 44-year-old female presenting with chest pain. "
    "She can be reached at 214-555-9876 or maria.gonzalez@email.com. "
    "Her IP address on the patient portal is 10.0.0.55. "
    "Vehicle on file: VIN 1HGBH41JXMN109186. "
    "Device serial: Serial #: DEV-2024-XYZ99. "
    "Referred by Dr. James Whitfield, License: Lic. TX-98765."
)

doc.add_heading("Insurance", level=1)
doc.add_paragraph("Medicare MBI: 1EG4-TE5-MK72 | Policy #: POL-2024-445566")

output_path = os.path.join("input_docs", "sample_patient_note.docx")
os.makedirs("input_docs", exist_ok=True)
doc.save(output_path)
print(f"Sample document created: {output_path}")