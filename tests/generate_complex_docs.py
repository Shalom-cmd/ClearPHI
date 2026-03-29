import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor
import json

OUTPUT_DIR = "input_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────
# DOC 1 — Hospital Discharge Summary
# ─────────────────────────────────────────────────────────────
def make_discharge_summary():
    doc = Document()
    doc.add_heading("DISCHARGE SUMMARY — CONFIDENTIAL", 0)

    doc.add_paragraph("Occidental Medical Center | 1200 W. Healthcare Blvd, Houston, TX 77001")
    doc.add_paragraph("Phone: (713) 555-0100 | Fax: Fax: (713) 555-0199 | www.occidentalmed.org")

    doc.add_heading("Patient Demographics", level=1)
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Full Name",           "Theodore James Harrington III"),
        ("Date of Birth",       "11/03/1959"),
        ("Age",                 "66 years old"),
        ("SSN",                 "321-54-9876"),
        ("MRN",                 "MRN#2048391"),
        ("Address",             "8821 Pinewood Drive, Apt 4B, Houston, TX 77002"),
        ("Phone",               "(713) 555-7823"),
        ("Emergency Contact",   "Patricia Harrington — (713) 555-7824"),
        ("Insurance",           "Member ID: BCB-TX-884421 | Policy #: POL-2024-119922"),
    ]
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_heading("Attending Physician", level=1)
    doc.add_paragraph(
        "Dr. Anita Krishnamurthy, MD | NPI: NPI:9876543210 | "
        "License: Lic. TX-MD-44821 | anita.k@occidentalmed.org"
    )

    doc.add_heading("Admission & Discharge Dates", level=1)
    doc.add_paragraph("Admitted: 03/10/2026 | Discharged: 03/18/2026")

    doc.add_heading("Clinical Notes", level=1)
    doc.add_paragraph(
        "Patient Theodore James Harrington III is a 66-year-old male admitted via ED "
        "on 03/10/2026 with acute onset chest pain radiating to the left arm. "
        "Troponin levels elevated at 2.4 ng/mL. EKG showed ST elevation in leads II, III, aVF. "
        "Patient denies tobacco use. Family history positive for CAD (father deceased at 71). "
        "Patient reached post-discharge at theodore.harrington@gmail.com or (713) 555-7823."
    )

    doc.add_heading("Medications at Discharge", level=1)
    meds = [
        "Metoprolol Succinate 50mg — once daily",
        "Aspirin 81mg — once daily",
        "Atorvastatin 40mg — once nightly",
        "Clopidogrel 75mg — once daily x 12 months",
    ]
    for med in meds:
        doc.add_paragraph(med, style="List Bullet")

    doc.add_heading("Follow-up Instructions", level=1)
    doc.add_paragraph(
        "Patient to follow up with Dr. Krishnamurthy within 7 days. "
        "Cardiology referral placed with Dr. Marcus Webb, NPI: NPI:1122334455. "
        "Patient portal: https://portal.occidentalmed.org/patient/TH2048391 "
        "Patient IP logged on portal: 192.168.10.45"
    )

    doc.add_heading("Device & Equipment", level=1)
    doc.add_paragraph(
        "Cardiac monitor serial: Serial #: CM-2026-XR9921 | "
        "Implantable loop recorder: Device ID: ILR-884421-B"
    )

    path = os.path.join(OUTPUT_DIR, "discharge_summary.docx")
    doc.save(path)
    print(f"✅ Created: {path}")


# ─────────────────────────────────────────────────────────────
# DOC 2 — Lab Report
# ─────────────────────────────────────────────────────────────
def make_lab_report():
    doc = Document()
    doc.add_heading("LABORATORY REPORT", 0)
    doc.add_paragraph("Occidental Diagnostics Lab | CLIA#: 45D2109876")
    doc.add_paragraph("Report Date: 03/14/2026 | Collected: 03/14/2026 07:30 AM")

    doc.add_heading("Patient", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Name",        "Rosa Mendez-Villarreal"),
        ("DOB",         "06/18/1988"),
        ("MRN",         "MRN#5519204"),
        ("SSN",         "789-45-1230"),
        ("Ordering MD", "Dr. Samuel Okonkwo | NPI: NPI:3344556677"),
    ]
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_heading("Results", level=1)
    results_table = doc.add_table(rows=6, cols=3)
    results_table.style = "Table Grid"
    headers = results_table.rows[0].cells
    headers[0].text = "Test"
    headers[1].text = "Result"
    headers[2].text = "Reference Range"

    results = [
        ("HbA1c",           "8.2%",         "< 5.7% normal"),
        ("Fasting Glucose",  "182 mg/dL",    "70-99 mg/dL"),
        ("LDL Cholesterol",  "134 mg/dL",    "< 100 mg/dL optimal"),
        ("eGFR",             "61 mL/min",    "> 60 mL/min normal"),
        ("TSH",              "2.1 mIU/L",    "0.4-4.0 mIU/L"),
    ]
    for i, (test, result, ref) in enumerate(results):
        row = results_table.rows[i+1]
        row.cells[0].text = test
        row.cells[1].text = result
        row.cells[2].text = ref

    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "Rosa Mendez-Villarreal, DOB 06/18/1988, presents with poorly controlled Type 2 DM. "
        "HbA1c of 8.2% indicates suboptimal glycemic control over past 3 months. "
        "LDL elevated — consider statin intensification. eGFR borderline — monitor renal function. "
        "Results faxed to ordering physician. Fax: (713) 555-0199. "
        "Patient notified via secure message at rosa.mv88@gmail.com."
    )

    path = os.path.join(OUTPUT_DIR, "lab_report.docx")
    doc.save(path)
    print(f"✅ Created: {path}")


# ─────────────────────────────────────────────────────────────
# DOC 3 — Referral Letter (tests abbreviations + shorthand)
# ─────────────────────────────────────────────────────────────
def make_referral_letter():
    doc = Document()
    doc.add_heading("REFERRAL LETTER", 0)

    doc.add_paragraph("From: Dr. Patricia Nguyen, MD | NPI: NPI:5566778899 | Lic. CA-MD-99201")
    doc.add_paragraph("Westside Family Practice | 400 Sunset Blvd, Los Angeles, CA 90028")
    doc.add_paragraph("Ph: (310) 555-2200 | Fax: Fax: (310) 555-2201 | pnguyen@westsideFP.com")
    doc.add_paragraph("Date: 03/20/2026")
    doc.add_paragraph("")

    doc.add_paragraph("To: Dr. Benjamin Okafor, MD — Nephrology")
    doc.add_paragraph("Cedars Specialty Group | 8700 Beverly Blvd, Los Angeles, CA 90048")
    doc.add_paragraph("")

    doc.add_heading("Re: Patient Referral", level=1)
    doc.add_paragraph(
        "Dear Dr. Okafor, I am referring my patient, Mr. Calvin J. Brooks, "
        "DOB 09/25/1951 (74 y/o male), MRN#: MRN#7734891, SSN: 654-32-1098 "
        "for nephrology evaluation. "
        "Pt c/o fatigue, LE edema x 3 wks. Labs notable for Cr 2.8, BUN 42, eGFR 24. "
        "Hx of HTN x 20yrs, T2DM — on metformin, lisinopril 20mg. "
        "Pt reachable at (310) 555-9988 or cjbrooks51@yahoo.com. "
        "Home address: 1122 Pacific Ave, Santa Monica, CA 90401. "
        "Insurance: Medicare MBI: 1EG4-TE5-MK72 | Acct#: Acct: 8844221199."
    )

    doc.add_paragraph(
        "Please note: pt has expressed preference for appts on Tues/Thurs. "
        "Interpreter needed — primary language Spanish. "
        "Vehicle on file for transport: VIN 2T1BURHE0JC987654. "
        "Thank you for seeing this patient."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Dr. Patricia Nguyen, MD")

    path = os.path.join(OUTPUT_DIR, "referral_letter.docx")
    doc.save(path)
    print(f"✅ Created: {path}")


# ─────────────────────────────────────────────────────────────
# DOC 4 — Edge case doc (tricky PHI patterns)
# ─────────────────────────────────────────────────────────────
def make_edge_cases():
    doc = Document()
    doc.add_heading("CLINICAL NOTE — EDGE CASE TEST", 0)
    doc.add_paragraph("This document contains intentionally tricky PHI patterns.")

    doc.add_heading("Age edge cases", level=1)
    doc.add_paragraph("Patient is a 34-year-old male.")
    doc.add_paragraph("Grandmother, age 92, also seen today.")
    doc.add_paragraph("Pediatric patient — 8 yr old female.")

    doc.add_heading("Date edge cases", level=1)
    doc.add_paragraph("Follow-up in 6 weeks (approx. late April 2026).")
    doc.add_paragraph("Surgery scheduled for 04/15/26.")
    doc.add_paragraph("Patient born in the winter of 1979.")

    doc.add_heading("Name edge cases", level=1)
    doc.add_paragraph("Referring MD: Dr. Kim (first name unknown).")
    doc.add_paragraph("Patient goes by 'Jay' — legal name: James Fitzgerald O'Brien.")
    doc.add_paragraph("POA: Ms. Loretta Washington-Banks, daughter.")

    doc.add_heading("Number edge cases", level=1)
    doc.add_paragraph("Dosage: 500mg — do not confuse with MRN.")
    doc.add_paragraph("Room 4021 — not a patient identifier.")
    doc.add_paragraph("BP: 142/88 — not a phone number.")
    doc.add_paragraph("Patient SSN on file: 112-23-3445.")
    doc.add_paragraph("ZIP of referring facility: 90028.")

    doc.add_heading("Mixed inline PHI", level=1)
    doc.add_paragraph(
        "Pt James Fitzgerald O'Brien (MRN#9920341, DOB 02/14/1968) was seen on 03/21/2026. "
        "Contact: jfobrien68@outlook.com | Cell: (424) 555-3344 | "
        "Alt contact: Loretta Washington-Banks (424) 555-3399. "
        "Employer: Pacific Logistics Inc, 500 Harbor Blvd, Long Beach CA 90802. "
        "Auth #: Member ID: PLN-2026-887733."
    )

    path = os.path.join(OUTPUT_DIR, "edge_cases.docx")
    doc.save(path)
    print(f"✅ Created: {path}")


# ─────────────────────────────────────────────────────────────
# RUN ALL
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\nGenerating synthetic patient documents...\n")
    make_discharge_summary()
    make_lab_report()
    make_referral_letter()
    make_edge_cases()
    print(f"\n✅ All documents saved to /{OUTPUT_DIR}/")
    print("Ready for de-identification testing.")